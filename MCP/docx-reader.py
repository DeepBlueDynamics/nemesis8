#!/usr/bin/env python3
"""
MCP: docx-reader

Utilities for reading DOCX files and extracting text in chunks.
"""

from __future__ import annotations

from pathlib import Path
from typing import Dict, List, Optional

from mcp.server.fastmcp import FastMCP

try:
    import docx  # python-docx
    DOCX_AVAILABLE = True
except Exception:
    docx = None
    DOCX_AVAILABLE = False

mcp = FastMCP("docx-reader")


def _load_doc(docx_path: Path):
    if not DOCX_AVAILABLE:
        raise RuntimeError("python-docx not installed in MCP venv.")
    return docx.Document(str(docx_path))


@mcp.tool()
async def docx_read_text(docx_path: str) -> Dict[str, object]:
    """
    Read all text from a DOCX file.
    """
    path = Path(docx_path)
    if not path.exists():
        return {"success": False, "error": "file_not_found", "path": str(path)}
    try:
        doc = _load_doc(path)
        paras = [p.text for p in doc.paragraphs if p.text]
        text = "\n".join(paras)
        return {
            "success": True,
            "path": str(path),
            "paragraphs": len(paras),
            "text": text,
        }
    except Exception as e:
        return {"success": False, "error": str(e), "path": str(path)}


@mcp.tool()
async def docx_list_paragraphs(docx_path: str) -> Dict[str, object]:
    """
    List paragraph texts with indices.
    """
    path = Path(docx_path)
    if not path.exists():
        return {"success": False, "error": "file_not_found", "path": str(path)}
    try:
        doc = _load_doc(path)
        paras = [{"index": i, "text": p.text} for i, p in enumerate(doc.paragraphs) if p.text]
        return {
            "success": True,
            "path": str(path),
            "paragraphs": paras,
            "count": len(paras),
        }
    except Exception as e:
        return {"success": False, "error": str(e), "path": str(path)}


@mcp.tool()
async def docx_extract_range(
    docx_path: str,
    start_index: int,
    end_index: int,
) -> Dict[str, object]:
    """
    Extract a range of paragraphs by index (inclusive).
    """
    path = Path(docx_path)
    if not path.exists():
        return {"success": False, "error": "file_not_found", "path": str(path)}
    try:
        doc = _load_doc(path)
        paras = [p.text for p in doc.paragraphs if p.text]
        if start_index < 0 or end_index < 0 or start_index > end_index:
            return {"success": False, "error": "invalid_range"}
        if start_index >= len(paras):
            return {"success": False, "error": "start_out_of_range", "count": len(paras)}
        end_index = min(end_index, len(paras) - 1)
        subset = paras[start_index:end_index + 1]
        return {
            "success": True,
            "path": str(path),
            "start_index": start_index,
            "end_index": end_index,
            "text": "\n".join(subset),
            "paragraph_count": len(subset),
        }
    except Exception as e:
        return {"success": False, "error": str(e), "path": str(path)}


if __name__ == "__main__":
    mcp.run()
